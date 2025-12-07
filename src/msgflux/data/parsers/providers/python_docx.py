from typing import Any, Dict, List, Optional, Union

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    Document = None
    Paragraph = None
    Table = None

from msgflux.data.parsers.base import BaseParser
from msgflux.data.parsers.registry import register_parser
from msgflux.data.parsers.response import ParserResponse
from msgflux.data.parsers.types import DocxParser
from msgflux.dotdict import dotdict


@register_parser
class PythonDocxDocxParser(BaseParser, DocxParser):
    """Python-DOCX-based DOCX Parser.

    Converts Word documents to Markdown format and extracts embedded images.
    Uses python-docx library for DOCX processing.

    Features:
    - Paragraph and heading extraction
    - Text formatting preservation
    - Table extraction (converted to Markdown)
    - Image extraction from document
    - Markdown output format

    Example:
        >>> parser = Parser.docx("python_docx")
        >>> response = parser("document.docx")
        >>> print(response.data["text"])
        >>> print(response.data["images"])
    """

    provider = "python_docx"

    def __init__(
        self,
        *,
        table_format: Optional[str] = "markdown",
        encode_images_base64: Optional[bool] = True,
    ):
        """Initialize Python-DOCX parser.

        Args:
            table_format:
                Format for table output. Options: "markdown" or "html".
            encode_images_base64:
                If True, encode images as base64 strings.
                If False, keep as raw bytes.
        """
        if Document is None:
            raise ImportError(
                "`python-docx` is not available. Install with `pip install python-docx`"
            )

        self.table_format = table_format
        self.encode_images_base64 = encode_images_base64
        self._initialize()

    def _initialize(self):
        """Initialize parser state."""
        pass

    def __call__(self, data: Union[str, bytes], **_kwargs) -> ParserResponse:
        """Parse a DOCX document.

        Args:
            data:
                DOCX file path, URL, or bytes.
            **kwargs:
                Additional parsing options (currently unused).

        Returns:
            ParserResponse containing:
            - text: Markdown-formatted content
            - images: Dictionary of extracted images
            - metadata: Document metadata (num_paragraphs, num_tables, etc.)

        Raises:
            FileNotFoundError:
                If file path doesn't exist.
            ValueError:
                If data type is not supported.
        """
        # Validate file type if it's a path
        if isinstance(data, str) and not data.startswith(("http://", "https://")):
            self._validate_file_type(data, ".docx")

        # Parse the document
        result = self._parse(data)

        # Create response
        response = ParserResponse()
        response.set_response_type("docx_parse")
        response.add(result)

        return response

    def _parse(self, data: bytes) -> List[Dict[str, Any]]:  # noqa: C901
        """Parse DOCX and extract content.

        Args:
            data:
                DOCX file path or bytes.

        Returns:
            Dictionary with:
            - text: Markdown content
            - images: Dictionary of images
            - metadata: Document metadata
        """
        md_content = ""
        images_dict = {}

        # Load document
        if isinstance(data, bytes):
            from io import BytesIO  # noqa: PLC0415

            document = Document(BytesIO(data))
        else:
            document = Document(data)

        # Extract images from document
        image_counter = 0
        for rel in document.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    img_data = rel.target_part.blob
                    # Get image extension from content type
                    content_type = rel.target_part.content_type
                    ext = content_type.split("/")[-1]
                    if ext == "jpeg":
                        ext = "jpg"

                    img_name = f"image_{image_counter}.{ext}"
                    images_dict[img_name] = img_data
                    image_counter += 1
                except Exception:  # noqa: S110
                    # Skip images that can't be extracted
                    pass

        # Process document elements (paragraphs and tables)
        for element in document.element.body:
            # Check if it's a paragraph
            if element.tag.endswith("p"):
                # Find the paragraph object
                for para in document.paragraphs:
                    if para._element == element:
                        para_text = self._process_paragraph(para)
                        if para_text:
                            md_content += para_text + "\n\n"
                        break

            # Check if it's a table
            elif element.tag.endswith("tbl"):
                # Find the table object
                for table in document.tables:
                    if table._element == element:
                        table_md = self._process_table(table)
                        if table_md:
                            md_content += table_md + "\n\n"
                        break

        # Prepare images (optionally encode to base64)
        images_dict = self._prepare_images_dict(
            images_dict, encode_base64=self.encode_images_base64
        )

        # Prepare metadata
        metadata = dotdict(
            {
                "num_paragraphs": len(document.paragraphs),
                "num_tables": len(document.tables),
                "num_images": len(images_dict),
                "table_format": self.table_format,
            }
        )

        # Try to extract document metadata if available
        if hasattr(document, "core_properties"):
            props = document.core_properties
            docx_metadata = {}
            if props.title:
                docx_metadata["title"] = props.title
            if props.author:
                docx_metadata["author"] = props.author
            if props.subject:
                docx_metadata["subject"] = props.subject
            if props.created:
                docx_metadata["created"] = str(props.created)
            if props.modified:
                docx_metadata["modified"] = str(props.modified)
            if docx_metadata:
                metadata.docx_info = docx_metadata

        return {
            "text": md_content.strip(),
            "images": images_dict,
            "metadata": metadata,
        }

    def _process_paragraph(self, paragraph) -> str:
        """Process a paragraph and convert to Markdown.

        Args:
            paragraph:
                Python-docx paragraph object.

        Returns:
            Markdown-formatted paragraph text.
        """
        text = paragraph.text.strip()
        if not text:
            return ""

        # Check paragraph style for headings
        style_name = paragraph.style.name.lower()

        if "heading 1" in style_name or style_name == "title":
            return f"# {text}"
        elif "heading 2" in style_name:
            return f"## {text}"
        elif "heading 3" in style_name:
            return f"### {text}"
        elif "heading 4" in style_name:
            return f"#### {text}"
        elif "heading 5" in style_name:
            return f"##### {text}"
        elif "heading 6" in style_name:
            return f"###### {text}"
        else:
            return text

    def _process_table(self, table) -> str:
        """Process a table and convert to Markdown or HTML.

        Args:
            table:
                Python-docx table object.

        Returns:
            Markdown or HTML formatted table.
        """
        # Extract table data
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)

        if not table_data:
            return ""

        # Convert to appropriate format
        if self.table_format == "html":
            return self._table_to_html(table_data)
        else:  # markdown
            return self._table_to_markdown(table_data)

    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """Convert table data to Markdown format.

        Args:
            table_data:
                2D list of table data.

        Returns:
            Markdown table string.
        """
        if not table_data:
            return ""

        # First row is header
        header = table_data[0]
        rows = table_data[1:]

        # Header line
        header_line = (
            "| " + " | ".join(str(cell) if cell else "" for cell in header) + " |"
        )

        # Separator line
        separator_line = "| " + " | ".join(["---"] * len(header)) + " |"

        # Data lines
        data_lines = [
            "| " + " | ".join(str(cell) if cell else "" for cell in row) + " |"
            for row in rows
        ]

        return "\n".join([header_line, separator_line, *data_lines])

    def _table_to_html(self, table_data: List[List[str]]) -> str:
        """Convert table data to HTML format.

        Args:
            table_data:
                2D list of table data.

        Returns:
            HTML table string.
        """
        if not table_data:
            return ""

        # First row is header
        header = table_data[0]
        rows = table_data[1:]

        # Start table
        html = "<table>\n"

        # Add header row
        html += "  <tr>\n"
        for cell in header:
            html += f"    <th>{str(cell) if cell else ''}</th>\n"
        html += "  </tr>\n"

        # Add data rows
        for row in rows:
            html += "  <tr>\n"
            for cell in row:
                html += f"    <td>{str(cell) if cell else ''}</td>\n"
            html += "  </tr>\n"

        # Close table
        html += "</table>"

        return html

    async def acall(self, data: Union[str, bytes], **_kwargs) -> ParserResponse:
        """Async version of __call__. Parse a DOCX document asynchronously.

        Args:
            data:
                DOCX file path, URL, or bytes.
            **kwargs:
                Additional parsing options (currently unused).

        Returns:
            ParserResponse containing:
            - text: Markdown-formatted content
            - images: Dictionary of extracted images
            - metadata: Document metadata (num_paragraphs, num_tables, etc.)

        Raises:
            FileNotFoundError:
                If file path doesn't exist.
            ValueError:
                If data type is not supported.
        """
        # Validate file type if it's a path
        if isinstance(data, str) and not data.startswith(("http://", "https://")):
            self._validate_file_type(data, ".docx")

        # Load file asynchronously if needed
        if isinstance(data, str):
            data = await self._aload_file(data)

        # Parse the document
        result = self._parse(data)

        # Create response
        response = ParserResponse()
        response.set_response_type("docx_parse")
        response.add(result)

        return response
